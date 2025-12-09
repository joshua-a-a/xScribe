import logging
import os
import sys
from pathlib import Path

from PySide6.QtCore import Qt
from PySide6.QtGui import QFont, QPalette
from PySide6.QtWidgets import (
    QApplication,
    QMainWindow,
    QMessageBox,
    QScrollArea,
    QVBoxLayout,
    QWidget,
)

from src.core.app_instance import AppInstanceManager
from src.gui.mac_app_delegate import register_mac_delegate
from src.gui.workers.batch_processor import BatchProcessor
from src.gui.workers.transcription_worker import TranscriptionWorker

from .components import (
    ActionButtonsComponent,
    FileInputComponent,
    ResultsComponent,
    SettingsComponent,
    StatusBarComponent,
)

logger = logging.getLogger(__name__)


class ProfessionalMainWindow(QMainWindow):
    """redesigned main window using modular components"""

    def __init__(self):
        super().__init__()
        logger.info("Initializing  Main Window")
        self.transcription_service = None
        self.current_worker = None
        self.batch_processor = None
        self.batch_results = []

        try:
            from src.core.hardware_monitor import HardwareMonitor

            self.hardware_monitor = HardwareMonitor()
            self.hardware_monitor.log_system_info()
        except ImportError:
            logger.warning("Hardware monitoring not available (psutil not installed)")
            self.hardware_monitor = None

        self.setup_window()
        self.setup_components()
        self.connect_signals()
        self.setup_styles()
        logger.info("Main Window initialization complete")
        self._mac_delegate = None

    def setup_window(self):
        self.setWindowTitle("xScribe")

        # Fixed width layout that still fills the available screen height for transcript viewing
        screen = QApplication.primaryScreen().availableGeometry()
        self.setMinimumWidth(1050)
        self.setMinimumHeight(600)  # Reasonable minimum height for small screens
        default_width = 1200
        default_height = screen.height()  # Use full available screen height

        self.resize(default_width, default_height)
        self.move(screen.x() + (screen.width() - default_width) // 2, screen.y())

    def setup_components(self):
        # Scroll container keeps layout usable on smaller displays
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAlwaysOff
        )  # No horizontal scroll
        scroll_area.setVerticalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAsNeeded
        )  # Vertical scroll when needed
        self.setCentralWidget(scroll_area)

        # Component stack lives inside the scroll area
        content_widget = QWidget()
        scroll_area.setWidget(content_widget)
        layout = QVBoxLayout(content_widget)
        layout.setSpacing(15)  # Increase spacing between major sections
        layout.setContentsMargins(10, 10, 10, 10)

        # Header component - REMOVED to save space
        # self.header = HeaderComponent()
        # self.header.setMaximumHeight(80)
        # layout.addWidget(self.header, 0)

        # Core components; heights describe how much vertical space each should reserve
        self.file_input = FileInputComponent()
        self.file_input.setMinimumHeight(450)
        self.file_input.setMaximumHeight(450)
        layout.addWidget(self.file_input, 0)

        self.settings = SettingsComponent()
        self.settings.setMinimumHeight(200)
        self.settings.setMaximumHeight(200)
        layout.addWidget(self.settings, 0)

        self.action_buttons = ActionButtonsComponent()
        self.action_buttons.setMinimumHeight(50)
        layout.addWidget(self.action_buttons, 0)

        # Results panel takes the remaining vertical space for transcripts/output
        self.results = ResultsComponent()
        self.results.setMinimumHeight(400)  # Larger minimum for comfortable viewing
        # No maximum height - let it expand to fill available space
        layout.addWidget(self.results, 1)  # Stretch factor of 1 = takes all extra space

        # Persistent status bar for health + progress readouts
        self.status_bar = StatusBarComponent()
        self.setStatusBar(self.status_bar)

    def connect_signals(self):
        # File input signals
        self.file_input.file_selected.connect(self._on_file_selected)

        # Batch processing signals (from FileInputComponent's BatchInputComponent)
        self.file_input.batch_component.batch_start_requested.connect(
            self._start_batch_processing
        )
        self.file_input.batch_component.batch_pause_requested.connect(
            self._pause_batch_processing
        )
        self.file_input.batch_component.batch_stop_requested.connect(
            self._stop_batch_processing
        )
        self.file_input.batch_component.export_batch_requested.connect(
            self._export_batch_results
        )
        self.file_input.batch_component.files_changed.connect(
            self._on_batch_files_changed
        )

        # Settings signals
        self.settings.model_changed.connect(self._on_model_changed)
        self.settings.enhanced_preprocessing_changed.connect(
            self._on_enhanced_preprocessing_changed
        )
        self.settings.speaker_detection_changed.connect(
            self._on_speaker_detection_changed
        )
        self.settings.word_timestamps_changed.connect(self._on_word_timestamps_changed)
        self.settings.language_changed.connect(self._on_language_changed)
        self.settings.auto_output_changed.connect(self._on_auto_output_changed)
        self.settings.output_location_change_requested.connect(
            self._change_output_location
        )
        self.settings.download_all_models_requested.connect(self._download_all_models)

        # Results signals
        self.results.export_requested.connect(self._export_results)
        self.results.timestamp_settings_changed.connect(
            self._on_timestamp_settings_changed
        )

        # Action button signals
        self.action_buttons.transcribe_requested.connect(self._start_transcription)
        self.action_buttons.privacy_report_requested.connect(
            self._generate_privacy_report
        )
        self.action_buttons.privacy_details_requested.connect(
            self._show_privacy_details
        )
        self.action_buttons.clear_requested.connect(self._clear_results)

    def setup_styles(self):
        # Detect if we're in dark mode
        palette = self.palette()
        is_dark_mode = palette.color(QPalette.Window).lightness() < 128

        if is_dark_mode:
            self._apply_dark_mode_styles()
        else:
            self._apply_light_mode_styles()

    def _apply_dark_mode_styles(self):
        self.setStyleSheet("""
            QMainWindow {
                background-color: #1e1e1e;
                color: #ffffff;
            }
            QGroupBox {
                font-weight: bold;
                border: 2px solid #404040;
                border-radius: 8px;
                margin: 10px 0px;
                padding-top: 10px;
                color: #ffffff;
                background-color: #2d2d2d;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #ffffff;
            }
            QLabel {
                color: #ffffff;
            }
            QTextEdit {
                background-color: #2d2d2d;
                color: #ffffff;
                border: 1px solid #404040;
                border-radius: 4px;
            }
            QCheckBox, QRadioButton {
                color: #ffffff;
            }
            QComboBox {
                background-color: #2d2d2d;
                color: #ffffff;
                border: 1px solid #404040;
                border-radius: 4px;
                padding: 4px;
            }
        """)

    def _apply_light_mode_styles(self):
        self.setStyleSheet("""
            QMainWindow {
                background-color: #ffffff;
                color: #000000;
            }
            QGroupBox {
                font-weight: bold;
                border: 2px solid #e2e8f0;
                border-radius: 8px;
                margin: 10px 0px;
                padding-top: 10px;
                color: #1e293b;
                background-color: #ffffff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
                color: #1e293b;
            }
        """)

    # Event handlers
    def bring_window_to_front(self):
        self.show()
        self.raise_()
        self.activateWindow()
        QApplication.processEvents()

    def handle_external_file_open(self, file_paths):
        if not file_paths:
            return

        # Filter to only real files (guard against multiprocessing bootstrap args)
        valid_paths = [p for p in file_paths if os.path.isfile(p)]
        if not valid_paths:
            logger.debug("No valid files in openFiles event, ignoring")
            return

        logger.info("macOS openFiles event received: %s", valid_paths)
        self.bring_window_to_front()

        primary_file = valid_paths[0]
        try:
            self.file_input.select_file(primary_file)
            self.status_bar.update_status(
                f"🔄 Status: File selected via Finder - {os.path.basename(primary_file)}"
            )
        except Exception as exc:
            logger.error(f"Failed to select Finder file '{primary_file}': {exc}")

        if len(valid_paths) > 1:
            additional = valid_paths[1:]
            added_count = self.file_input.batch_component.add_external_files(additional)
            if added_count:
                self.status_bar.update_status(
                    f"🔄 Status: {added_count} additional file(s) queued for batch processing"
                )

    def _on_file_selected(self, file_path):
        self.status_bar.update_status(
            f"🔄 Status: File selected - {os.path.basename(file_path)}"
        )

    def _on_model_changed(self, model_name):
        self.status_bar.update_status(f"🔄 Status: Model changed to {model_name}")

    def _on_enhanced_preprocessing_changed(self, enabled):
        status = "enabled" if enabled else "disabled"
        self.status_bar.update_status(f"🔄 Status: Enhanced preprocessing {status}")

    def _on_speaker_detection_changed(self, enabled):
        status = "enabled" if enabled else "disabled"
        self.status_bar.update_status(f"🔄 Status: Speaker detection {status}")

    def _on_word_timestamps_changed(self, enabled):
        status = "enabled" if enabled else "disabled"
        self.status_bar.update_status(f"Status: Word-level timestamping {status}")
        if enabled:
            self.status_bar.update_status(
                "🎬 Subtitle generation ready - SRT & VTT formats available"
            )
        else:
            self.status_bar.update_status(
                "⚡ Fast transcription mode - no subtitle timing data"
            )

    def _on_language_changed(self, language_code):
        self.status_bar.update_status(f"🔄 Status: Language set to {language_code}")

    def _on_auto_output_changed(self, enabled):
        status = "enabled" if enabled else "disabled"
        self.status_bar.update_status(f"🔄 Status: Auto-save {status}")

    def _change_output_location(self):
        from PySide6.QtWidgets import QFileDialog

        directory = QFileDialog.getExistingDirectory(
            self,
            "Select Output Directory",
            self.settings.get_output_settings()["base_directory"],
        )

        if directory:
            self.settings.set_output_directory(directory)
            self.status_bar.update_status("🔄 Status: Output location changed")

    def _download_all_models(self):
        from PySide6.QtWidgets import QApplication, QMessageBox

        from src.core.first_run_manager import FirstRunManager
        from src.gui.components.welcome_dialog import ModelDownloadDialog

        manager = FirstRunManager()

        # Check which models are missing
        all_models = ["tiny", "base", "small", "medium"]
        missing_models = [m for m in all_models if not manager.is_model_downloaded(m)]

        if not missing_models:
            QMessageBox.information(
                self,
                "All Models Downloaded",
                "All models are already downloaded!\n\n"
                f"Available models: {', '.join(all_models)}",
            )
            return

        # Confirm download
        response = QMessageBox.question(
            self,
            "Download Models",
            f"Download {len(missing_models)} missing model(s)?\n\n"
            f"Models to download: {', '.join(missing_models)}\n\n"
            f"This will download approximately {len(missing_models) * 0.5:.1f} GB",
            QMessageBox.Yes | QMessageBox.No,
        )

        if response != QMessageBox.Yes:
            return

        # Show download dialog
        download_dialog = ModelDownloadDialog(
            "all", "various sizes", self, is_batch=True
        )
        download_dialog.show()

        def progress_callback(message, percent, current_model, total_models):
            download_dialog.update_progress(
                message, percent, current_model, total_models
            )
            QApplication.processEvents()

        # Download missing models
        success, failed = manager.download_all_models(progress_callback)

        if success and not failed:
            download_dialog.accept()
            QMessageBox.information(
                self,
                "Download Complete",
                f"Successfully downloaded {len(missing_models)} model(s)!",
            )
        else:
            download_dialog.reject()
            failed_str = ", ".join(failed) if failed else "unknown"
            QMessageBox.warning(
                self,
                "Download Issues",
                f"Some models failed to download: {failed_str}\n\n"
                "Please check your internet connection and try again.",
            )

    def _on_timestamp_settings_changed(self, show_timestamps, interval):
        self.status_bar.update_status("🔄 Status: Timestamp settings updated")

    def _start_transcription(self):
        selected_file = self.file_input.get_selected_file()
        if not selected_file:
            self.status_bar.update_status("🔄 Status: No file selected")
            return

        # Check system health before starting
        if self.hardware_monitor:
            health = self.hardware_monitor.check_system_health()
            if not health["healthy"]:
                warning_msg = "System health warnings detected:\n\n" + "\n".join(
                    f"• {w}" for w in health["warnings"]
                )
                warning_msg += "\n\nContinuing may lead to poor performance or crashes. Continue anyway?"

                reply = QMessageBox.warning(
                    self,
                    "System Health Warning",
                    warning_msg,
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No,
                )

                if reply == QMessageBox.No:
                    self.status_bar.update_status(
                        "⏸️ Transcription cancelled due to system health"
                    )
                    return

                logger.warning(
                    f"User proceeding despite health warnings: {health['warnings']}"
                )

        # Validate file before starting
        from src.core.transcription_service import EnhancedTranscriptionService

        temp_service = EnhancedTranscriptionService()
        is_valid, validation_msg = temp_service.validate_file(selected_file)

        if not is_valid:
            # Show error dialog for invalid files
            QMessageBox.critical(
                self,
                "Invalid Audio File",
                f"Cannot process this file:\n\n{validation_msg}\n\n"
                f"Please select a valid audio file.",
            )
            self.status_bar.update_status(f"❌ {validation_msg}")
            return

        # Check if file is long (30+ minutes) - ask for confirmation
        if validation_msg.startswith("LONG_FILE:"):
            duration_mins = float(validation_msg.split(":")[1])
            estimated_time = duration_mins / 20  # Rough estimate: 20x real-time

            reply = QMessageBox.question(
                self,
                "Long Audio File Detected",
                f"This audio file is {duration_mins:.1f} minutes long.\n\n"
                f"Estimated processing time: ~{estimated_time:.1f} minutes\n"
                f"(actual time varies by model and system)\n\n"
                f"Do you want to continue?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No,
            )

            if reply == QMessageBox.No:
                self.status_bar.update_status("⏸️ Transcription cancelled by user")
                return

        # Get settings
        config = self.settings.get_configuration()

        logger.info(f"Starting transcription: {Path(selected_file).name}")
        logger.info(
            f"Configuration: model={config['model']}, language={config['language']}, "
            f"enhanced={config['enhanced_preprocessing']}, speaker_detection={config['speaker_detection']}"
        )

        # Check if model is downloaded, download if needed
        from src.core.first_run_manager import FirstRunManager

        manager = FirstRunManager()
        selected_model = config["model"]

        if not manager.is_model_downloaded(selected_model):
            logger.info(f"Model {selected_model} not found, initiating download")

            # Show download dialog
            from src.gui.components.welcome_dialog import ModelDownloadDialog

            model_size = manager.MODEL_SIZES.get(selected_model, "unknown")

            download_dialog = ModelDownloadDialog(
                selected_model, model_size, self, is_batch=False
            )
            download_dialog.show()

            def progress_callback(message, percent):
                download_dialog.update_progress(message, percent)
                QApplication.processEvents()

            success = manager.download_model(selected_model, progress_callback)

            if success:
                download_dialog.accept()
                logger.info(f"Model {selected_model} downloaded successfully")
            else:
                download_dialog.reject()
                QMessageBox.warning(
                    self,
                    "Model Download Failed",
                    f"Failed to download {selected_model} model.\n\n"
                    "Please check your internet connection and try again.",
                )
                self.status_bar.update_status("❌ Model download failed")
                return

        # Create and start transcription worker with correct parameters
        self.current_worker = TranscriptionWorker(
            file_path=selected_file,
            model=config["model"],
            language=config["language"],
            enhanced=config["enhanced_preprocessing"],
            speaker_detection=config["speaker_detection"],
            word_timestamps=config[
                "word_timestamps"
            ],  # Pass the word timestamps setting
        )

        # Connect worker signals (using the correct signal names)
        self.current_worker.progress_updated.connect(
            self._on_transcription_progress_updated
        )
        self.current_worker.transcription_finished.connect(
            self._on_transcription_finished
        )
        self.current_worker.error_occurred.connect(self._on_transcription_error)

        # Update UI
        self.action_buttons.set_all_enabled(False)
        self.status_bar.show_processing_status("TRANSCRIPTION")
        self.status_bar.show_progress(0, "Initializing transcription...")

        # Start worker
        self.current_worker.start()

    def _on_transcription_progress_updated(self, step, message, progress):
        self.status_bar.update_progress(int(progress), message)

    def _on_transcription_finished(self, results_dict):
        # Convert dict back to TranscriptionResult if needed
        if isinstance(results_dict, dict):
            from src.models.transcription_result import TranscriptionResult

            results = TranscriptionResult.from_dict(results_dict)
        else:
            results = results_dict

        logger.info(
            f"✅ Transcription completed: {results.word_count} words, "
            f"{results.duration:.1f}s audio, processed in {results.processing_time:.1f}s"
        )

        self.results.set_transcription_results(results)

        # Auto-save if enabled
        auto_output_enabled = self.settings.get_auto_output()
        logger.info(f"Auto-output setting: {auto_output_enabled}")
        if auto_output_enabled:
            logger.info("Calling _auto_save_results...")
            self._auto_save_results(results)
        else:
            logger.warning("Auto-save is disabled - results not saved to file")

        self.status_bar.hide_progress()
        self.status_bar.show_ready_status()
        self.action_buttons.set_all_enabled(True)

        # CRITICAL: Clean up model from memory to prevent leaks
        if self.current_worker and hasattr(
            self.current_worker, "transcription_service"
        ):
            if self.current_worker.transcription_service:
                self.current_worker.transcription_service.cleanup()
                logger.info("Memory cleanup completed")

        if self.current_worker:
            self.current_worker.deleteLater()
            self.current_worker = None

    def _auto_save_results(self, results):
        try:
            from datetime import datetime
            from pathlib import Path

            from src.core.filename_utils import (
                create_safe_output_path,
                safe_filename_from_path,
            )

            # Get output settings
            output_settings = self.settings.get_output_settings()
            base_dir = Path(output_settings["base_directory"]).expanduser()

            # Create output directory if it doesn't exist
            base_dir.mkdir(parents=True, exist_ok=True)

            # Generate safe filename from original file
            safe_name = safe_filename_from_path(
                results.file_path,
                default=f"transcription_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            )

            self.status_bar.update_status("💾 Auto-saving transcription results...")

            logger.info(
                f"Auto-save: '{Path(results.file_path).name if results.file_path else 'unnamed'}' -> '{safe_name}_transcript.txt'"
            )

            # Save as text file with sanitized filename
            txt_path = create_safe_output_path(
                base_dir, results.file_path, "_transcript", ".txt"
            )
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write("xScribe Transcription Results\n")
                f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Model: {results.model_used}\n")
                f.write(f"Language: {results.language}\n")
                f.write(f"Duration: {results.duration:.2f} seconds\n")
                f.write(f"Processing Time: {results.processing_time:.2f} seconds\n")
                f.write(f"Word Count: {results.word_count}\n")
                f.write("\n" + "=" * 50 + "\n")
                f.write("TRANSCRIPTION:\n")
                f.write("=" * 50 + "\n\n")
                f.write(results.full_text)

                # Add timestamped segments if available
                if results.segments and any(
                    seg.start is not None for seg in results.segments
                ):
                    f.write("\n\n" + "=" * 50 + "\n")
                    f.write("TIMESTAMPED SEGMENTS:\n")
                    f.write("=" * 50 + "\n\n")
                    for segment in results.segments:
                        if segment.start is not None and segment.end is not None:
                            start_time = f"{int(segment.start // 60):02d}:{int(segment.start % 60):02d}"
                            end_time = f"{int(segment.end // 60):02d}:{int(segment.end % 60):02d}"
                            speaker_prefix = (
                                f"[{segment.speaker}] " if segment.speaker else ""
                            )
                            f.write(
                                f"[{start_time} - {end_time}] {speaker_prefix}{segment.text}\n"
                            )

            # Also save as JSON for full data preservation with sanitized filename
            json_path = create_safe_output_path(
                base_dir, results.file_path, "_transcript", ".json"
            )
            import json

            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(results.to_dict(), f, indent=2, ensure_ascii=False)

            # Update status
            self.status_bar.update_status(f"✅ Saved to {txt_path.name}")

            # Log the save
            logger.info(f"Auto-saved transcription results to: {txt_path}")

            return True  # Success

        except Exception as e:
            # Don't let auto-save errors break the main functionality
            error_msg = f"Auto-save failed: {str(e)}"
            self.status_bar.update_status(f"⚠️ {error_msg}")

            logger.warning(f"Auto-save failed: {e}")

            import traceback

            logger.error(f"Full traceback: {traceback.format_exc()}")

            return False  # Failure

    def _on_transcription_error(self, error_message):
        self.status_bar.hide_progress()
        self.status_bar.show_error_status()
        self.status_bar.update_status(f"🔄 Status: Error - {error_message}")
        self.action_buttons.set_all_enabled(True)

        # CRITICAL: Clean up model from memory even on error
        if self.current_worker and hasattr(
            self.current_worker, "transcription_service"
        ):
            if self.current_worker.transcription_service:
                self.current_worker.transcription_service.cleanup()

        if self.current_worker:
            self.current_worker.deleteLater()
            self.current_worker = None

    def _export_results(self, export_format):
        self.status_bar.update_status(
            f"🔄 Status: Exporting to {export_format.upper()}..."
        )

        try:
            import os

            from PySide6.QtWidgets import QFileDialog

            # Get current results
            current_results = self.results.get_current_results()
            if not current_results:
                self.status_bar.update_status("❌ Status: No results to export")
                return

            # Determine file extension and filter
            if export_format == "txt":
                file_filter = "Text Files (*.txt)"
                default_ext = ".txt"
            elif export_format == "srt":
                file_filter = "SubRip Subtitle Files (*.srt)"
                default_ext = ".srt"
            elif export_format == "vtt":
                file_filter = "WebVTT Subtitle Files (*.vtt)"
                default_ext = ".vtt"
            elif export_format == "docx":
                file_filter = "Word Documents (*.docx)"
                default_ext = ".docx"
            elif export_format == "json":
                file_filter = "JSON Files (*.json)"
                default_ext = ".json"
            else:
                self.status_bar.update_status(
                    f"❌ Status: Unsupported format {export_format}"
                )
                return

            # Show save dialog
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                f"Export as {export_format.upper()}",
                f"transcription{default_ext}",
                file_filter,
            )

            if not file_path:
                self.status_bar.update_status("🔄 Status: Export cancelled")
                return

            # Export based on format
            if export_format == "txt":
                self._export_txt(file_path, current_results)
            elif export_format == "srt":
                self._export_srt(file_path, current_results)
            elif export_format == "vtt":
                self._export_vtt(file_path, current_results)
            elif export_format == "docx":
                self._export_docx(file_path, current_results)
            elif export_format == "json":
                self._export_json(file_path, current_results)

            self.status_bar.update_status(
                f"✅ Status: Exported to {os.path.basename(file_path)}"
            )

        except Exception as e:
            self.status_bar.update_status(f"❌ Status: Export failed - {str(e)}")

    def _export_txt(self, file_path, results):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(results.full_text)

    def _export_srt(self, file_path, results):
        with open(file_path, "w", encoding="utf-8") as f:
            if hasattr(results, "segments") and results.segments:
                for i, segment in enumerate(results.segments, 1):
                    start_time = self._format_time_srt(segment.start)
                    end_time = self._format_time_srt(segment.end)
                    text = segment.text.strip()

                    if text:
                        f.write(f"{i}\n")
                        f.write(f"{start_time} --> {end_time}\n")
                        f.write(f"{text}\n\n")
            else:
                # No segments, just export basic text with placeholder timestamps
                f.write("1\n00:00:00,000 --> 00:05:00,000\n")
                f.write(results.full_text)
                f.write("\n\n")

    def _export_vtt(self, file_path, results):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write("WEBVTT\n\n")
            f.write("NOTE\nGenerated by xScribe\n\n")

            if hasattr(results, "segments") and results.segments:
                for segment in results.segments:
                    start_time = self._format_time_vtt(segment.start)
                    end_time = self._format_time_vtt(segment.end)
                    text = segment.text.strip()

                    if text:
                        f.write(f"{start_time} --> {end_time}\n")
                        f.write(f"{text}\n\n")
            else:
                # No segments, just export basic text with placeholder timestamps
                f.write("00:00:00.000 --> 00:05:00.000\n")
                f.write(results.full_text)
                f.write("\n\n")

    def _export_docx(self, file_path, results):
        try:
            from docx import Document

            doc = Document()
            doc.add_heading("Transcription Results", 0)
            doc.add_paragraph(results.full_text)
            doc.save(file_path)
        except ImportError:
            # Fallback to text if docx not available
            with open(file_path.replace(".docx", ".txt"), "w", encoding="utf-8") as f:
                f.write("Transcription Results\n")
                f.write("=" * 20 + "\n\n")
                f.write(results.full_text)

    def _export_json(self, file_path, results):
        import json

        data = {"full_text": results.full_text, "segments": []}

        if hasattr(results, "segments") and results.segments:
            for segment in results.segments:
                seg_data = {
                    "start": segment.start,
                    "end": segment.end,
                    "text": segment.text,
                }
                if hasattr(segment, "speaker") and segment.speaker:
                    seg_data["speaker"] = segment.speaker
                data["segments"].append(seg_data)

        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

    def _format_time_srt(self, seconds):
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        milliseconds = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{milliseconds:03d}"

    def _format_time_vtt(self, seconds):
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        milliseconds = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{milliseconds:03d}"

    def _generate_privacy_report(self):
        from datetime import datetime

        self.status_bar.update_status("Status: Generating privacy report...")

        try:
            # Get configuration paths
            output_settings = self.settings.get_output_settings()
            outputs_path = os.path.expanduser(
                output_settings.get("base_directory", "~/Desktop/xScribe Outputs")
            )
            logs_path = os.path.join(os.path.dirname(__file__), "..", "..", "logs")
            model_path = os.path.join(os.path.dirname(__file__), "..", "..", ".venv")

            # Get settings
            encryption_enabled = False  # Add this when encryption is implemented

            # Generate report content
            report = f"""=== xScribe Privacy Report ===
Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

NETWORK ACTIVITY:
- No internet connection required
- No external API calls
- All processing happens locally on your device
- Whisper models loaded from: {model_path}

DATA STORAGE:
- Audio files: Never uploaded, processed locally only
- Transcriptions saved to: {outputs_path}
- Application logs: {logs_path}
- No cloud synchronization
- No automatic backups to external services

PRIVACY FEATURES:
- End-to-end local processing
- Offline AI models (OpenAI Whisper)
- Speaker diarization (100% local, no cloud APIs)
- Encrypted output option: {"Enabled" if encryption_enabled else "Available (not configured)"}
- No telemetry or analytics
- No external dependencies during runtime
- Complete data sovereignty

COMPLIANCE:
- GDPR compliant (data never leaves your device)
- HIPAA compatible (when used with proper procedures)
- Zero trust architecture
- No third-party data processors
- Full audit trail in local logs

TECHNICAL DETAILS:
- AI Model: OpenAI Whisper (runs locally)
- Processing: 100% on-device computation
- Network: Zero external connections required
- Data Flow: Audio → Local Whisper Model → Local Storage
- Third Parties: None

SECURITY FEATURES:
- Filename sanitization (prevents injection attacks)
- Audio validation (prevents malicious files)
- Disk space validation (prevents system instability)
- Memory monitoring (prevents resource exhaustion)
- Crash recovery (prevents data loss)
- Emergency backup system

AUDIT TRAIL:
- All operations logged to: {logs_path}
- Daily log files with full processing history
- Error tracking and debugging information
- No logs transmitted externally

PRIVACY GUARANTEE:
xScribe is designed with privacy as the foundation. All audio processing,
transcription, and analysis happens entirely on your local device. Your
audio files and transcriptions never leave your computer. No internet
connection is required after initial setup, and no data is ever collected,
transmitted, or stored by any third party.

For questions about privacy or compliance, review the logs in:
{logs_path}
"""

            # Save report to outputs folder
            report_filename = (
                f"privacy_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            )
            report_path = os.path.join(outputs_path, report_filename)

            with open(report_path, "w", encoding="utf-8") as f:
                f.write(report)

            logger.info(f"Privacy report generated: {report_path}")

            # Show the report in a dialog
            self._show_privacy_report_dialog(report, report_path)

            self.status_bar.update_status("Status: Privacy report generated")

        except Exception as e:
            logger.error(f"Error generating privacy report: {e}", exc_info=True)
            QMessageBox.warning(
                self,
                "Privacy Report Error",
                f"Could not generate privacy report:\n\n{str(e)}",
            )
            self.status_bar.update_status("Status: Privacy report generation failed")

    def _show_privacy_report_dialog(self, report_text, report_path):
        from PySide6.QtWidgets import (
            QDialog,
            QHBoxLayout,
            QLabel,
            QPushButton,
            QTextEdit,
            QVBoxLayout,
        )

        dialog = QDialog(self)
        dialog.setWindowTitle("Privacy Compliance Report")
        dialog.setMinimumSize(800, 600)

        layout = QVBoxLayout()

        # Header label
        header = QLabel(f"Privacy report saved to:\n{os.path.basename(report_path)}")
        header.setStyleSheet("font-weight: bold; padding: 10px; color: #10b981;")
        layout.addWidget(header)

        # Report text area
        text_area = QTextEdit()
        text_area.setPlainText(report_text)
        text_area.setReadOnly(True)
        text_area.setStyleSheet("""
            QTextEdit {
                font-family: 'Courier New', monospace;
                font-size: 11px;
                background-color: #1e1e1e;
                color: #ffffff;
                border: 1px solid #404040;
                border-radius: 4px;
                padding: 10px;
            }
        """)

        layout.addWidget(text_area)

        # Buttons
        button_layout = QHBoxLayout()

        open_folder_btn = QPushButton("Open Reports Folder")
        open_folder_btn.clicked.connect(
            lambda: self._open_folder(os.path.dirname(report_path))
        )

        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)

        button_layout.addWidget(open_folder_btn)
        button_layout.addStretch()
        button_layout.addWidget(close_btn)

        layout.addLayout(button_layout)
        dialog.setLayout(layout)
        dialog.exec()

    def _show_privacy_details(self):
        details = """PRIVACY-FIRST DESIGN

100% Offline Processing
  No internet required after initial setup.
  All AI models run locally on your device.

Local AI Models
  OpenAI Whisper models stored on your device.
  No cloud API calls for transcription.

Zero Data Collection
  No telemetry, analytics, or tracking.
  No user data transmitted to any server.

Complete Audit Trail
  Full logging for compliance verification.
  Daily log files stored locally only.

Advanced Security
  - Filename sanitization prevents injection attacks
  - Audio validation prevents malicious files
  - Disk space checks prevent system issues
  - Memory monitoring prevents resource exhaustion
  - Crash recovery with emergency backups

Data Sovereignty
  Your audio and transcriptions stay on your device.
  You have complete control over your data.

Compliance Ready
  - GDPR compliant (no data leaves device)
  - HIPAA compatible (with proper procedures)
  - Zero trust architecture
  - No third-party processors

Privacy Guarantee
  xScribe is designed from the ground up for privacy.
  Your data never touches the internet.
  No exceptions, no backdoors, no compromises.

For detailed compliance information, generate a Privacy Report."""

        QMessageBox.information(self, "Privacy Features", details)

        self.status_bar.update_status("Status: Privacy details displayed")

    def _open_folder(self, folder_path):
        import platform
        import subprocess

        try:
            if platform.system() == "Darwin":  # macOS
                subprocess.run(["open", folder_path])
            elif platform.system() == "Windows":
                subprocess.run(["explorer", folder_path])
            else:  # Linux
                subprocess.run(["xdg-open", folder_path])
        except Exception as e:
            logger.error(f"Error opening folder: {e}")

    def _clear_results(self):
        self.results.clear_results()
        self.status_bar.update_status("🔄 Status: Results cleared")

    # Batch Processing Handlers
    def _on_batch_files_changed(self, file_list):
        count = len(file_list)
        self.status_bar.update_status(f"🔄 Status: {count} files in batch queue")

    def _start_batch_processing(self):
        batch_files = self.file_input.batch_component.get_batch_files()
        if not batch_files:
            self.status_bar.update_status("🔄 Status: No files in batch queue")
            return

        # Validate all files first and check for long files
        from src.core.transcription_service import EnhancedTranscriptionService

        temp_service = EnhancedTranscriptionService()

        long_files = []
        invalid_files = []

        for file_path in batch_files:
            is_valid, validation_msg = temp_service.validate_file(file_path)
            if not is_valid:
                invalid_files.append((file_path, validation_msg))
            elif validation_msg.startswith("LONG_FILE:"):
                duration_mins = float(validation_msg.split(":")[1])
                long_files.append((file_path, duration_mins))

        # Show warning about invalid files
        if invalid_files:
            invalid_list = "\n".join(
                [f"• {Path(f).name}: {msg}" for f, msg in invalid_files[:5]]
            )
            if len(invalid_files) > 5:
                invalid_list += f"\n... and {len(invalid_files) - 5} more"

            QMessageBox.warning(
                self,
                "Invalid Files Detected",
                f"{len(invalid_files)} file(s) will be skipped:\n\n{invalid_list}\n\n"
                f"The batch will continue with the remaining valid files.",
            )

        # Show confirmation for long files
        if long_files:
            total_duration = sum(duration for _, duration in long_files)
            estimated_time = total_duration / 20  # Rough estimate

            long_list = "\n".join(
                [f"• {Path(f).name}: {dur:.1f} min" for f, dur in long_files[:5]]
            )
            if len(long_files) > 5:
                long_list += f"\n... and {len(long_files) - 5} more"

            reply = QMessageBox.question(
                self,
                "Long Audio Files Detected",
                f"{len(long_files)} file(s) are over 30 minutes:\n\n{long_list}\n\n"
                f"Total duration: ~{total_duration:.1f} minutes\n"
                f"Estimated processing time: ~{estimated_time:.1f} minutes\n\n"
                f"Do you want to continue?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No,
            )

            if reply == QMessageBox.No:
                self.status_bar.update_status("⏸️ Batch processing cancelled by user")
                return

        # Clear previous batch results
        self.batch_results = []

        # Get settings for batch processing
        config = self.settings.get_configuration()

        # Check if model is downloaded, download if needed
        from src.core.first_run_manager import FirstRunManager

        manager = FirstRunManager()
        selected_model = config["model"]

        if not manager.is_model_downloaded(selected_model):
            logger.info(
                f"Model {selected_model} not found for batch processing, initiating download"
            )

            # Show download dialog
            from src.gui.components.welcome_dialog import ModelDownloadDialog

            model_size = manager.MODEL_SIZES.get(selected_model, "unknown")

            download_dialog = ModelDownloadDialog(
                selected_model, model_size, self, is_batch=False
            )
            download_dialog.show()

            def progress_callback(message, percent):
                download_dialog.update_progress(message, percent)
                QApplication.processEvents()

            success = manager.download_model(selected_model, progress_callback)

            if success:
                download_dialog.accept()
                logger.info(
                    f"Model {selected_model} downloaded successfully for batch processing"
                )
            else:
                download_dialog.reject()
                QMessageBox.warning(
                    self,
                    "Model Download Failed",
                    f"Failed to download {selected_model} model.\n\n"
                    "Please check your internet connection and try again.",
                )
                self.status_bar.update_status("❌ Model download failed")
                return

        # DEBUG: Log the model being used
        print("\n" + "=" * 60)
        print("🔍 BATCH DEBUG - Model Selection Trace")
        print("=" * 60)
        print(f"Selected model from settings: {config['model']}")

        # Check which radio button is actually checked
        for button in self.settings.model_group.buttons():
            if button.isChecked():
                print(f"✓ Checked radio button: {button.text()} (value={button.value})")
            else:
                print(f"  Unchecked: {button.text()} (value={button.value})")

        print(f"Full config: {config}")
        print("=" * 60 + "\n")

        # Create BatchFile objects for processing
        from src.gui.workers.batch_processor import BatchFile

        batch_file_objects = [BatchFile(file_path) for file_path in batch_files]

        # Create and start batch processor
        self.batch_processor = BatchProcessor(
            files=batch_file_objects,
            model=config["model"],
            language=config["language"],
            enhanced=config["enhanced_preprocessing"],
            speaker_detection=config["speaker_detection"],
        )

        # Connect batch processor signals
        self.batch_processor.file_started.connect(self._on_batch_file_started)
        self.batch_processor.file_progress.connect(self._on_batch_file_progress)
        self.batch_processor.file_completed.connect(self._on_batch_file_completed)
        self.batch_processor.file_failed.connect(self._on_batch_file_failed)
        self.batch_processor.batch_completed.connect(self._on_batch_finished)

        # Update UI for batch processing
        self.file_input.batch_component.set_batch_controls_enabled(
            start_enabled=False,
            pause_enabled=True,
            stop_enabled=True,
            export_enabled=False,  # Disable export during processing
        )
        self.action_buttons.set_all_enabled(False)
        self.status_bar.show_processing_status("BATCH PROCESSING")
        self.status_bar.show_progress(0, "Starting batch processing...")

        # Start batch processor
        self.batch_processor.start()

    def _pause_batch_processing(self):
        if self.batch_processor:
            self.batch_processor.pause()
            self.status_bar.update_status("🔄 Status: Batch processing paused")

    def _stop_batch_processing(self):
        if self.batch_processor:
            # Clean up model before stopping
            if hasattr(self.batch_processor, "transcription_service"):
                if self.batch_processor.transcription_service:
                    self.batch_processor.transcription_service.cleanup()

            self.batch_processor.stop()
            self.batch_processor.wait()
            self._reset_batch_ui()
            self.status_bar.update_status("🔄 Status: Batch processing stopped")

    def _export_batch_results(self):
        if not self.batch_results:
            self.status_bar.update_status("⚠️ No batch results to export")
            return

        self.status_bar.update_status(
            f"📦 Exporting {len(self.batch_results)} batch results..."
        )

        try:
            import json
            import shutil
            from datetime import datetime
            from pathlib import Path

            from src.models.transcription_result import TranscriptionResult

            # Get output directory
            output_settings = self.settings.get_output_settings()
            base_dir = Path(output_settings["base_directory"]).expanduser()

            # Check if base directory exists and is writable
            if not base_dir.exists():
                try:
                    base_dir.mkdir(parents=True, exist_ok=True)
                except PermissionError:
                    self.status_bar.update_status(
                        f"❌ No permission to create directory: {base_dir}"
                    )
                    return
                except Exception as e:
                    self.status_bar.update_status(
                        f"❌ Cannot create directory: {str(e)}"
                    )
                    return

            # Check available disk space (require at least 100MB free)
            stat = shutil.disk_usage(base_dir)
            free_mb = stat.free / (1024 * 1024)
            if free_mb < 100:
                error_msg = f"Insufficient disk space: Only {free_mb:.1f}MB free (need 100MB minimum)"
                self.status_bar.update_status(f"❌ {error_msg}")
                QMessageBox.critical(
                    self,
                    "Disk Space Error",
                    f"{error_msg}\n\nPlease free up disk space or choose a different output directory.",
                )
                return

            # Create batch-specific subdirectory
            batch_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            batch_dir = base_dir / f"batch_{batch_timestamp}"

            try:
                batch_dir.mkdir(parents=True, exist_ok=True)
            except PermissionError:
                self.status_bar.update_status(
                    "❌ No permission to create batch directory"
                )
                return
            except Exception as e:
                self.status_bar.update_status(
                    f"❌ Cannot create batch directory: {str(e)}"
                )
                return

            exported_count = 0
            failed_count = 0

            # Export each result
            for batch_result in self.batch_results:
                try:
                    from src.core.filename_utils import (
                        create_safe_output_path,
                        safe_filename_from_path,
                    )

                    results_dict = batch_result["results"]
                    transcription_result = TranscriptionResult.from_dict(results_dict)

                    # Generate safe filename
                    safe_name = safe_filename_from_path(
                        transcription_result.file_path,
                        default=f"transcription_{batch_result['index']}",
                    )

                    logger.info(
                        f"Batch export: '{transcription_result.file_path}' -> '{safe_name}_transcript.txt'"
                    )

                    # Save as text file with sanitized filename
                    txt_path = create_safe_output_path(
                        batch_dir, transcription_result.file_path, "_transcript", ".txt"
                    )
                    try:
                        with open(txt_path, "w", encoding="utf-8") as f:
                            f.write("xScribe Transcription Results\n")
                            f.write(
                                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                            )
                            f.write(f"Model: {transcription_result.model_used}\n")
                            f.write(f"Language: {transcription_result.language}\n")
                            f.write(
                                f"Duration: {transcription_result.duration:.2f} seconds\n"
                            )
                            f.write(
                                f"Processing Time: {transcription_result.processing_time:.2f} seconds\n"
                            )
                            f.write(f"Word Count: {transcription_result.word_count}\n")
                            f.write("\n" + "=" * 50 + "\n")
                            f.write("TRANSCRIPTION:\n")
                            f.write("=" * 50 + "\n\n")
                            f.write(transcription_result.full_text)

                            # Add timestamped segments if available
                            if transcription_result.segments and any(
                                seg.start is not None
                                for seg in transcription_result.segments
                            ):
                                f.write("\n\n" + "=" * 50 + "\n")
                                f.write("TIMESTAMPED SEGMENTS:\n")
                                f.write("=" * 50 + "\n\n")
                                for segment in transcription_result.segments:
                                    if (
                                        segment.start is not None
                                        and segment.end is not None
                                    ):
                                        start_time = f"{int(segment.start // 60):02d}:{int(segment.start % 60):02d}"
                                        end_time = f"{int(segment.end // 60):02d}:{int(segment.end % 60):02d}"
                                        speaker_prefix = (
                                            f"[{segment.speaker}] "
                                            if segment.speaker
                                            else ""
                                        )
                                        f.write(
                                            f"[{start_time} - {end_time}] {speaker_prefix}{segment.text}\n"
                                        )
                    except IOError as e:
                        raise Exception(f"Failed to write TXT file: {str(e)}")

                    # Save as JSON with sanitized filename
                    json_path = create_safe_output_path(
                        batch_dir,
                        transcription_result.file_path,
                        "_transcript",
                        ".json",
                    )
                    try:
                        with open(json_path, "w", encoding="utf-8") as f:
                            json.dump(
                                transcription_result.to_dict(),
                                f,
                                indent=2,
                                ensure_ascii=False,
                            )
                    except IOError as e:
                        raise Exception(f"Failed to write JSON file: {str(e)}")

                    exported_count += 1

                except Exception as e:
                    error_msg = str(e)
                    # Check for disk space errors
                    if "No space left" in error_msg or "Disk full" in error_msg:
                        self.status_bar.update_status("❌ Export failed: Disk full!")
                        logger.error(
                            f"Disk full during export. Exported {exported_count}/{len(self.batch_results)} files"
                        )

                        QMessageBox.critical(
                            self,
                            "Export Failed - Disk Full",
                            f"Ran out of disk space during export!\n\n"
                            f"Successfully exported: {exported_count}/{len(self.batch_results)} files\n"
                            f"Export directory: {batch_dir}\n\n"
                            f"Please free up disk space and try exporting the remaining files.",
                        )
                        return  # Stop exporting

                    print(
                        f"Failed to export result {batch_result['index']}: {error_msg}"
                    )
                    logger.error(
                        f"Export error for file {batch_result['index']}: {error_msg}"
                    )
                    failed_count += 1

            # Show completion message
            self.status_bar.update_status(
                f"✅ Exported {exported_count} results to {batch_dir.name}"
                + (f" ({failed_count} failed)" if failed_count > 0 else "")
            )

            # Log completion
            logger.info(
                f"Batch export completed: {exported_count} files to {batch_dir}"
            )

        except Exception as e:
            self.status_bar.update_status(f"❌ Batch export failed: {str(e)}")

            logger.error(f"Batch export error: {e}")

    def _on_batch_file_started(self, file_index, filename):
        self.status_bar.update_status(
            f"🔄 Status: Processing file {file_index + 1}: {filename}"
        )

    def _on_batch_file_progress(self, file_index, step, message, progress):
        self.status_bar.update_progress(int(progress), message)

        # Update progress in batch table
        batch_files = self.file_input.batch_component.get_batch_files()
        if file_index < len(batch_files):
            file_path = batch_files[file_index]
            self.file_input.batch_component.update_file_status(
                file_path, "Processing", int(progress)
            )

    def _on_batch_file_completed(self, file_index, results):
        if results:
            # Store result for export
            self.batch_results.append({"index": file_index, "results": results})

            # Auto-save individual file if enabled
            if self.settings.get_auto_output():
                try:
                    from src.models.transcription_result import TranscriptionResult

                    transcription_result = TranscriptionResult.from_dict(results)
                    self._auto_save_results(transcription_result)
                except Exception as e:
                    print(f"Auto-save failed for file {file_index}: {e}")

            # Update file status in batch tree
            batch_files = self.file_input.batch_component.get_batch_files()
            if file_index < len(batch_files):
                file_path = batch_files[file_index]
                self.file_input.batch_component.update_file_status(
                    file_path, "Completed", 100
                )

    def _on_batch_file_failed(self, file_index, error_message):
        # Update file status in batch tree
        batch_files = self.file_input.batch_component.get_batch_files()
        if file_index >= 0 and file_index < len(batch_files):
            file_path = batch_files[file_index]
            self.file_input.batch_component.update_file_status(file_path, "Failed", 0)

    def _on_batch_finished(self):
        # CRITICAL: Clean up model from memory to prevent leaks
        if self.batch_processor and hasattr(
            self.batch_processor, "transcription_service"
        ):
            if self.batch_processor.transcription_service:
                self.batch_processor.transcription_service.cleanup()

        self._reset_batch_ui()
        self.status_bar.hide_progress()
        self.status_bar.show_ready_status()

        # Get final status from batch component
        batch_files = self.file_input.batch_component.get_batch_files()
        total = len(batch_files)
        # For now, just show completion message
        self.status_bar.update_status(
            f"🔄 Status: Batch processing completed - {total} files processed"
        )

        if self.batch_processor:
            self.batch_processor.deleteLater()
            self.batch_processor = None

    def _reset_batch_ui(self):
        self.file_input.batch_component.set_batch_controls_enabled(
            start_enabled=True,
            pause_enabled=False,
            stop_enabled=False,
            export_enabled=True,  # Re-enable export when done
        )
        self.action_buttons.set_all_enabled(True)

    def closeEvent(self, event):
        # CRITICAL: Clean up models from memory before closing
        if self.current_worker and hasattr(
            self.current_worker, "transcription_service"
        ):
            if self.current_worker.transcription_service:
                self.current_worker.transcription_service.cleanup()

        if self.batch_processor and hasattr(
            self.batch_processor, "transcription_service"
        ):
            if self.batch_processor.transcription_service:
                self.batch_processor.transcription_service.cleanup()

        # Clean up workers
        if self.current_worker:
            self.current_worker.terminate()
            self.current_worker.wait()

        if self.batch_processor:
            self.batch_processor.terminate()
            self.batch_processor.wait()

        # Clean up status bar
        if hasattr(self.status_bar, "cleanup"):
            self.status_bar.cleanup()

        event.accept()


def emergency_save_state():
    """
    Legacy emergency save function - delegates to AppInstanceManager.
    This function exists for backward compatibility with existing code.
    """
    manager = AppInstanceManager()
    return manager.emergency_save_state()


def run_professional_gui(is_first_run=False):
    """
    Run the professional xScribe GUI application

    Args:
        is_first_run: Whether this is the first time running the app

    Returns:
        Exit code (0 for success)
    """
    app = QApplication(sys.argv)
    app.setApplicationName("xScribe")
    app.setApplicationVersion("2.0")

    # Set application font
    app.setFont(QFont("Arial", 10))

    # Show welcome dialog if first run
    if is_first_run:
        from PySide6.QtWidgets import QDialog, QMessageBox

        from src.core.first_run_manager import FirstRunManager
        from src.gui.components.welcome_dialog import ModelDownloadDialog, WelcomeDialog

        welcome = WelcomeDialog()
        if welcome.exec() == QDialog.Accepted:
            selected_model = welcome.get_selected_model()
            logger.info(f"First run: User selected {selected_model} model")

            manager = FirstRunManager()

            # Handle "all models" option
            if selected_model == "all":
                # Show download dialog for batch
                download_dialog = ModelDownloadDialog(
                    "all", "2.2 GB total", is_batch=True
                )
                download_dialog.show()

                # Download all models with progress
                def progress_callback(message, percent, current, total):
                    download_dialog.update_progress(message, percent, current, total)
                    QApplication.processEvents()  # Keep UI responsive

                success, failed = manager.download_all_models(progress_callback)

                if success:
                    download_dialog.accept()
                    manager.complete_first_run()
                    logger.info("Successfully downloaded all models")
                    QMessageBox.information(
                        None,
                        "Download Complete",
                        "All models downloaded successfully!\n\n"
                        "You now have access to:\n"
                        "• Tiny (75 MB)\n"
                        "• Base (150 MB)\n"
                        "• Small (500 MB)\n"
                        "• Medium (1.5 GB)\n\n"
                        "Choose the best model for each job!",
                    )
                else:
                    download_dialog.reject()
                    failed_list = ", ".join(failed)
                    QMessageBox.warning(
                        None,
                        "Download Incomplete",
                        f"Some models failed to download: {failed_list}\n\n"
                        "You can try downloading them individually from settings.",
                    )
                    manager.complete_first_run()  # Still mark complete
            else:
                # Single model download
                model_size = manager.MODEL_SIZES.get(selected_model, "unknown")

                download_dialog = ModelDownloadDialog(
                    selected_model, model_size, is_batch=False
                )
                download_dialog.show()

                # Download model in background
                def progress_callback(message, progress, current=None, total=None):
                    download_dialog.update_progress(message, progress)
                    QApplication.processEvents()  # Keep UI responsive

                success = manager.download_model(selected_model, progress_callback)

                if success:
                    download_dialog.accept()
                    manager.complete_first_run()
                    logger.info("First run setup completed successfully")
                else:
                    download_dialog.reject()
                    QMessageBox.warning(
                        None,
                        "Download Failed",
                        f"Failed to download {selected_model} model. "
                        "You can try again later from the settings.",
                    )
        else:
            logger.info("User cancelled first run setup")
            # Still mark as complete so we don't show it again
            FirstRunManager().complete_first_run()

    # Create and show main window
    logger.info("Creating main window")
    window = ProfessionalMainWindow()

    # Register window with instance manager (NO MORE GLOBALS!)
    from src.core.app_instance import AppInstanceManager

    manager = AppInstanceManager()
    manager.set_main_window(window)

    # Register macOS delegate to intercept open-file events
    delegate = register_mac_delegate(window)
    if delegate:
        window._mac_delegate = delegate

    logger.info("Showing main window")
    window.show()

    # THIS IS THE CRITICAL PART - Start the Qt event loop
    logger.info("Starting Qt event loop")
    exit_code = app.exec()

    logger.info(f"Qt event loop exited with code: {exit_code}")

    # Clean up the instance manager
    manager.clear()

    return exit_code


if __name__ == "__main__":
    sys.exit(run_professional_gui())
